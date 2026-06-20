#!/usr/bin/env python3
"""
RAG CHUNKER -- Unified document chunking engine (S99).

Provides a single entry point for chunking any supported document type
(PDF, DOCX, XLSX, CSV, plain text, code) with configurable parameters.

Each chunk carries rich metadata: source_file, page/section, absolute
position within the document, and a parent_doc_id for traceability.

Delegates to the existing opti_oignon.rag.chunkers infrastructure where
possible, adding PDF/DOCX/XLSX text extraction and overlap injection.
"""

import hashlib
import logging
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# -- Optional heavy imports (graceful degradation) -------------------------

PYPDF_AVAILABLE = False
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PdfReader = None  # type: ignore[assignment,misc]

DOCX_AVAILABLE = False
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None  # type: ignore[assignment,misc]

OPENPYXL_AVAILABLE = False
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    openpyxl = None  # type: ignore[assignment]

# Feature flag for downstream guards
RAG_CHUNKER_AVAILABLE = True


# =========================================================================
# DATA STRUCTURES
# =========================================================================

@dataclass
class RAGChunk:
    """A chunk produced by RAGChunker with full provenance metadata."""

    content: str
    source_file: str
    file_type: str
    chunk_index: int
    total_chunks: int
    parent_doc_id: str
    position: int = 0          # Absolute character offset in source text
    page: int | None = None    # Page number (PDF only)
    section: str | None = None # Section or heading name
    start_line: int | None = None
    end_line: int | None = None

    @property
    def chunk_id(self) -> str:
        """Deterministic unique ID for this chunk."""
        raw = f"{self.parent_doc_id}::{self.chunk_index}"
        return hashlib.sha256(raw.encode()).hexdigest()[:16]

    @property
    def metadata(self) -> dict[str, Any]:
        """Flat metadata dict suitable for ChromaDB storage."""
        return {
            "source_file": self.source_file,
            "file_type": self.file_type,
            "chunk_index": self.chunk_index,
            "total_chunks": self.total_chunks,
            "parent_doc_id": self.parent_doc_id,
            "position": self.position,
            "page": self.page if self.page is not None else -1,
            "section": self.section or "",
            "start_line": self.start_line if self.start_line is not None else -1,
            "end_line": self.end_line if self.end_line is not None else -1,
            "char_count": len(self.content),
        }

    def __repr__(self) -> str:
        preview = self.content[:60].replace("\n", " ")
        return (
            f"RAGChunk(idx={self.chunk_index}/{self.total_chunks}, "
            f"file={Path(self.source_file).name}, "
            f"section={self.section!r}, "
            f"chars={len(self.content)}, "
            f"preview={preview!r}...)"
        )


@dataclass
class ChunkingResult:
    """Result of a chunking operation on a single document."""

    doc_id: str
    source_file: str
    file_type: str
    chunks: list[RAGChunk]
    raw_text_length: int
    extraction_method: str  # "pdf", "docx", "xlsx", "csv", "text", "direct"

    @property
    def chunk_count(self) -> int:
        return len(self.chunks)


# =========================================================================
# CHUNKER
# =========================================================================

class RAGChunker:
    """
    Unified document chunking engine.

    Supports PDF, DOCX, XLSX/XLS, CSV/TSV, Markdown, plain text,
    and code files (Python, R, JS, TS, etc.).

    Usage::

        chunker = RAGChunker(chunk_size=500, chunk_overlap=50)
        result = chunker.chunk_file("/path/to/doc.pdf")
        for chunk in result.chunks:
            print(chunk.chunk_id, chunk.section, len(chunk.content))

        # Or from raw text:
        result = chunker.chunk_text("Hello world...", source="inline.txt")
    """

    # Extension -> file_type mapping
    EXT_MAP: dict[str, str] = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".xlsx": "excel",
        ".xls": "excel",
        ".csv": "csv",
        ".tsv": "csv",
        ".md": "markdown",
        ".markdown": "markdown",
        ".txt": "text",
        ".text": "text",
        ".py": "python",
        ".r": "r",
        ".R": "r",
        ".rmd": "rmarkdown",
        ".Rmd": "rmarkdown",
        ".js": "javascript",
        ".ts": "typescript",
        ".html": "html",
        ".css": "css",
        ".json": "json",
        ".yaml": "yaml",
        ".yml": "yaml",
        ".toml": "toml",
        ".sh": "shell",
        ".sql": "sql",
    }

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        """
        Args:
            chunk_size: Target chunk size in approximate tokens.
            chunk_overlap: Number of overlap characters between adjacent chunks.
        """
        self.chunk_size = max(50, chunk_size)
        self.chunk_overlap = max(0, min(chunk_overlap, chunk_size // 2))

    # -----------------------------------------------------------------
    # PUBLIC API
    # -----------------------------------------------------------------

    def chunk_file(
        self,
        filepath: str | Path,
        doc_id: str | None = None,
    ) -> ChunkingResult:
        """
        Chunk a file from disk.

        Handles text extraction for binary formats (PDF, DOCX, XLSX)
        and delegates to the appropriate chunking strategy.

        Args:
            filepath: Path to the file.
            doc_id: Optional document ID (generated if absent).

        Returns:
            ChunkingResult with all chunks and metadata.
        """
        filepath = Path(filepath)
        if not filepath.exists():
            raise FileNotFoundError(f"File not found: {filepath}")

        doc_id = doc_id or uuid.uuid4().hex[:12]
        file_type = self.EXT_MAP.get(filepath.suffix.lower(), "text")

        # Extract text from binary formats
        text, extraction = self._extract_text(filepath, file_type)
        if not text or len(text.strip()) < 10:
            return ChunkingResult(
                doc_id=doc_id,
                source_file=str(filepath),
                file_type=file_type,
                chunks=[],
                raw_text_length=len(text) if text else 0,
                extraction_method=extraction,
            )

        chunks = self._do_chunk(text, str(filepath), file_type, doc_id)

        return ChunkingResult(
            doc_id=doc_id,
            source_file=str(filepath),
            file_type=file_type,
            chunks=chunks,
            raw_text_length=len(text),
            extraction_method=extraction,
        )

    def chunk_text(
        self,
        text: str,
        source: str = "inline",
        file_type: str = "text",
        doc_id: str | None = None,
    ) -> ChunkingResult:
        """
        Chunk raw text content.

        Args:
            text: The text to chunk.
            source: A label for the source (filename, URL, etc.).
            file_type: The file type hint for choosing a chunking strategy.
            doc_id: Optional document ID.

        Returns:
            ChunkingResult with all chunks.
        """
        doc_id = doc_id or uuid.uuid4().hex[:12]

        if not text or len(text.strip()) < 10:
            return ChunkingResult(
                doc_id=doc_id,
                source_file=source,
                file_type=file_type,
                chunks=[],
                raw_text_length=len(text) if text else 0,
                extraction_method="direct",
            )

        chunks = self._do_chunk(text, source, file_type, doc_id)

        return ChunkingResult(
            doc_id=doc_id,
            source_file=source,
            file_type=file_type,
            chunks=chunks,
            raw_text_length=len(text),
            extraction_method="direct",
        )

    # -----------------------------------------------------------------
    # TEXT EXTRACTION (binary formats)
    # -----------------------------------------------------------------

    def _extract_text(
        self, filepath: Path, file_type: str
    ) -> tuple[str, str]:
        """
        Extract text from a file.

        Returns (text, extraction_method).
        """
        if file_type == "pdf":
            return self._extract_pdf(filepath), "pdf"
        elif file_type == "docx":
            return self._extract_docx(filepath), "docx"
        elif file_type == "excel":
            return self._extract_excel(filepath), "xlsx"
        else:
            # Text-based: read directly
            try:
                text = filepath.read_text(encoding="utf-8", errors="replace")
                return text, "text"
            except Exception as exc:
                logger.error("Failed to read %s: %s", filepath, exc)
                return "", "text"

    @staticmethod
    def _extract_pdf(filepath: Path) -> str:
        """Extract text from PDF using pypdf."""
        if not PYPDF_AVAILABLE:
            logger.error("pypdf not installed. Install with: pip install pypdf")
            return ""
        try:
            reader = PdfReader(str(filepath))
            parts: list[str] = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    parts.append(f"[Page {i + 1}]\n{page_text}")
            return "\n\n".join(parts)
        except Exception as exc:
            logger.error("PDF extraction failed for %s: %s", filepath, exc)
            return ""

    @staticmethod
    def _extract_docx(filepath: Path) -> str:
        """Extract text from DOCX preserving heading hierarchy."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return ""
        try:
            doc = DocxDocument(str(filepath))
            parts: list[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = (para.style.name.lower() if para.style else "")
                if "heading 1" in style_name or "titre 1" in style_name:
                    parts.append(f"\n# {text}")
                elif "heading 2" in style_name or "titre 2" in style_name:
                    parts.append(f"\n## {text}")
                elif "heading 3" in style_name or "titre 3" in style_name:
                    parts.append(f"\n### {text}")
                elif "title" in style_name:
                    parts.append(f"\n# {text}")
                else:
                    parts.append(text)
            # Tables
            for table in doc.tables:
                rows_text: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        rows_text.append(" | ".join(cells))
                if rows_text:
                    parts.append("\n[Table]\n" + "\n".join(rows_text))
            return "\n\n".join(parts)
        except Exception as exc:
            logger.error("DOCX extraction failed for %s: %s", filepath, exc)
            return ""

    @staticmethod
    def _extract_excel(filepath: Path) -> str:
        """Extract text from Excel workbook."""
        if not OPENPYXL_AVAILABLE:
            logger.error("openpyxl not installed. Install with: pip install openpyxl")
            return ""
        try:
            wb = openpyxl.load_workbook(str(filepath), data_only=True)
            parts: list[str] = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows: list[str] = []
                for row in sheet.iter_rows(values_only=True):
                    row_str = [str(c) if c is not None else "" for c in row]
                    if any(s.strip() for s in row_str):
                        rows.append(",".join(row_str))
                if rows:
                    parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
            return "\n\n".join(parts)
        except Exception as exc:
            logger.error("Excel extraction failed for %s: %s", filepath, exc)
            return ""

    # -----------------------------------------------------------------
    # CHUNKING STRATEGIES
    # -----------------------------------------------------------------

    def _do_chunk(
        self,
        text: str,
        source_file: str,
        file_type: str,
        doc_id: str,
    ) -> list[RAGChunk]:
        """
        Route to the appropriate chunking strategy and produce RAGChunks.
        """
        # Detect page markers for PDF content
        page_map = self._build_page_map(text) if file_type == "pdf" else {}

        if file_type == "pdf":
            raw_chunks = self._chunk_pdf_text(text)
        elif file_type in ("markdown", "rmarkdown"):
            raw_chunks = self._chunk_markdown(text)
        elif file_type in ("csv",):
            raw_chunks = self._chunk_csv(text)
        elif file_type in ("excel",):
            raw_chunks = self._chunk_csv(text)  # Already text with headers
        elif file_type in ("python", "javascript", "typescript", "shell", "sql", "r"):
            raw_chunks = self._chunk_code(text)
        elif file_type == "docx":
            raw_chunks = self._chunk_markdown(text)  # DOCX extracted as markdown-like
        else:
            raw_chunks = self._chunk_plain_text(text)

        # Apply overlap between adjacent chunks
        raw_chunks = self._apply_overlap(raw_chunks, text)

        # Build RAGChunk objects
        chunks: list[RAGChunk] = []
        total = len(raw_chunks)
        for i, (content, position, section, start_line, end_line) in enumerate(raw_chunks):
            page = self._find_page(position, page_map) if page_map else None
            chunks.append(RAGChunk(
                content=content,
                source_file=source_file,
                file_type=file_type,
                chunk_index=i,
                total_chunks=total,
                parent_doc_id=doc_id,
                position=position,
                page=page,
                section=section,
                start_line=start_line,
                end_line=end_line,
            ))

        return chunks

    # -- PDF chunking (by page/paragraph) --

    def _chunk_pdf_text(self, text: str) -> list[tuple[str, int, str | None, int | None, int | None]]:
        """Chunk PDF text that contains [Page N] markers."""
        return self._chunk_plain_text(text)

    # -- Markdown / DOCX chunking (by heading) --

    def _chunk_markdown(self, text: str) -> list[tuple[str, int, str | None, int | None, int | None]]:
        """Chunk by markdown headings."""
        header_re = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        headers = list(header_re.finditer(text))

        if not headers:
            return self._chunk_plain_text(text)

        max_chars = self.chunk_size * 4
        raw: list[tuple[str, int, str | None, int | None, int | None]] = []

        # Preamble before first header
        if headers[0].start() > 0:
            preamble = text[:headers[0].start()].strip()
            if len(preamble) >= 10:
                raw.append((preamble, 0, "preamble", 1, text[:headers[0].start()].count("\n") + 1))

        for i, match in enumerate(headers):
            level = len(match.group(1))
            title = match.group(2).strip()
            start_pos = match.start()

            end_pos = len(text)
            for j in range(i + 1, len(headers)):
                if len(headers[j].group(1)) <= level:
                    end_pos = headers[j].start()
                    break

            section_text = text[start_pos:end_pos].strip()
            if len(section_text) < 10:
                continue

            start_line = text[:start_pos].count("\n") + 1
            end_line = start_line + section_text.count("\n")

            # Subdivide if too large
            if len(section_text) > max_chars:
                sub_parts = self._recursive_split(section_text, max_chars)
                pos = start_pos
                for sp in sub_parts:
                    sl = text[:pos].count("\n") + 1
                    el = sl + sp.count("\n")
                    raw.append((sp, pos, f"h{level}: {title}", sl, el))
                    pos += len(sp)
            else:
                raw.append((section_text, start_pos, f"h{level}: {title}", start_line, end_line))

        return raw

    # -- CSV / Excel chunking (row-based with headers) --

    def _chunk_csv(self, text: str) -> list[tuple[str, int, str | None, int | None, int | None]]:
        """Chunk CSV keeping header row with each chunk."""
        lines = text.strip().split("\n")
        if len(lines) < 2:
            return [(text, 0, "data", 1, len(lines))]

        header = lines[0]
        data_lines = lines[1:]

        max_chars = self.chunk_size * 4
        avg_len = sum(len(l) for l in data_lines[:10]) // max(1, min(10, len(data_lines)))
        rows_per_chunk = max(5, max_chars // max(1, avg_len + 1))

        raw: list[tuple[str, int, str | None, int | None, int | None]] = []
        pos = 0
        for i in range(0, len(data_lines), rows_per_chunk):
            batch = data_lines[i : i + rows_per_chunk]
            content = header + "\n" + "\n".join(batch)
            section = f"rows_{i + 1}_to_{i + len(batch)}"
            start_line = i + 2
            end_line = i + len(batch) + 1
            raw.append((content, pos, section, start_line, end_line))
            pos += sum(len(l) + 1 for l in batch)

        return raw

    # -- Code chunking --

    _FUNC_RE = re.compile(r'^(?:async\s+)?def\s+(\w+)\s*\(', re.MULTILINE)
    _CLASS_RE = re.compile(r'^class\s+(\w+)\s*[:(]', re.MULTILINE)

    def _chunk_code(self, text: str) -> list[tuple[str, int, str | None, int | None, int | None]]:
        """Chunk code by function/class boundaries."""
        boundaries: list[tuple[int, str]] = []

        for m in self._CLASS_RE.finditer(text):
            boundaries.append((m.start(), f"class: {m.group(1)}"))
        for m in self._FUNC_RE.finditer(text):
            line_start = text.rfind("\n", 0, m.start()) + 1
            if m.start() - line_start == 0:  # Top-level only
                boundaries.append((m.start(), f"function: {m.group(1)}"))

        if not boundaries:
            return self._chunk_plain_text(text)

        boundaries.sort(key=lambda b: b[0])
        max_chars = self.chunk_size * 4
        raw: list[tuple[str, int, str | None, int | None, int | None]] = []

        for i, (start, name) in enumerate(boundaries):
            end = boundaries[i + 1][0] if i + 1 < len(boundaries) else len(text)
            chunk_text = text[start:end].strip()
            if len(chunk_text) < 10:
                continue
            sl = text[:start].count("\n") + 1
            el = sl + chunk_text.count("\n")
            if len(chunk_text) > max_chars:
                sub = self._recursive_split(chunk_text, max_chars)
                pos = start
                for s in sub:
                    raw.append((s, pos, name, text[:pos].count("\n") + 1, text[:pos].count("\n") + 1 + s.count("\n")))
                    pos += len(s)
            else:
                raw.append((chunk_text, start, name, sl, el))

        return raw if raw else self._chunk_plain_text(text)

    # -- Plain text (recursive character splitting) --

    def _chunk_plain_text(self, text: str) -> list[tuple[str, int, str | None, int | None, int | None]]:
        """Recursive character splitting for plain text."""
        max_chars = self.chunk_size * 4
        parts = self._recursive_split(text, max_chars)

        raw: list[tuple[str, int, str | None, int | None, int | None]] = []
        pos = 0
        for i, part in enumerate(parts):
            sl = text[:pos].count("\n") + 1
            el = sl + part.count("\n")
            raw.append((part, pos, f"segment_{i + 1}", sl, el))
            # Advance position (find next occurrence to handle overlap trimming)
            idx = text.find(part[:40], pos) if len(part) >= 40 else pos
            pos = idx + len(part) if idx >= 0 else pos + len(part)

        return raw

    def _recursive_split(self, text: str, max_chars: int) -> list[str]:
        """Split text recursively using progressively finer separators."""
        if len(text) <= max_chars:
            return [text] if text.strip() else []

        separators = ["\n\n\n", "\n\n", "\n", ". ", " "]
        for sep in separators:
            parts = text.split(sep)
            if len(parts) > 1:
                merged = self._merge_small_parts(parts, sep, max_chars)
                if all(len(m) <= max_chars * 1.2 for m in merged):
                    return merged

        # Hard split as last resort
        result: list[str] = []
        for i in range(0, len(text), max_chars):
            chunk = text[i : i + max_chars].strip()
            if chunk:
                result.append(chunk)
        return result

    @staticmethod
    def _merge_small_parts(parts: list[str], sep: str, max_chars: int) -> list[str]:
        """Merge small parts back together up to max_chars."""
        merged: list[str] = []
        current = ""
        for part in parts:
            candidate = current + sep + part if current else part
            if len(candidate) <= max_chars:
                current = candidate
            else:
                if current.strip():
                    merged.append(current.strip())
                current = part
        if current.strip():
            merged.append(current.strip())
        return merged

    # -----------------------------------------------------------------
    # OVERLAP INJECTION
    # -----------------------------------------------------------------

    def _apply_overlap(
        self,
        raw_chunks: list[tuple[str, int, str | None, int | None, int | None]],
        full_text: str,
    ) -> list[tuple[str, int, str | None, int | None, int | None]]:
        """Inject overlap from adjacent chunks."""
        if self.chunk_overlap <= 0 or len(raw_chunks) <= 1:
            return raw_chunks

        ov = self.chunk_overlap
        result: list[tuple[str, int, str | None, int | None, int | None]] = []
        for i, (content, pos, section, sl, el) in enumerate(raw_chunks):
            prefix = ""
            suffix = ""
            if i > 0:
                prev_content = raw_chunks[i - 1][0]
                prefix = prev_content[-ov:]
            if i < len(raw_chunks) - 1:
                next_content = raw_chunks[i + 1][0]
                suffix = next_content[:ov]

            parts: list[str] = []
            if prefix:
                parts.append(prefix)
            parts.append(content)
            if suffix:
                parts.append(suffix)

            merged = "\n...\n".join(parts) if len(parts) > 1 else content
            result.append((merged, pos, section, sl, el))

        return result

    # -----------------------------------------------------------------
    # PAGE MAP (PDF)
    # -----------------------------------------------------------------

    @staticmethod
    def _build_page_map(text: str) -> dict[int, int]:
        """
        Build a position -> page number map from [Page N] markers.

        Returns dict mapping character offset to page number.
        """
        page_re = re.compile(r'\[Page (\d+)\]')
        page_map: dict[int, int] = {}
        for m in page_re.finditer(text):
            page_map[m.start()] = int(m.group(1))
        return page_map

    @staticmethod
    def _find_page(position: int, page_map: dict[int, int]) -> int | None:
        """Find which page a given character position belongs to."""
        if not page_map:
            return None
        best_page = None
        for marker_pos, page_num in sorted(page_map.items()):
            if marker_pos <= position:
                best_page = page_num
            else:
                break
        return best_page


# =========================================================================
# MODULE-LEVEL SINGLETON
# =========================================================================

_default_chunker: RAGChunker | None = None


def get_rag_chunker(
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> RAGChunker:
    """Return the module-level RAGChunker singleton."""
    global _default_chunker
    if _default_chunker is None:
        _default_chunker = RAGChunker(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
    return _default_chunker
