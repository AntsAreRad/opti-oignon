#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CHUNKERS - Découpage intelligent des documents
==============================================
Stratégies de chunking adaptées à chaque type de fichier.

Le chunking respecte la structure logique des documents :
- Code : par fonction/classe/section
- Markdown : par headers
- Texte : par paragraphes
"""

import re
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple
import logging

from .config import get_config, ChunkingConfig

logger = logging.getLogger(__name__)


@dataclass
class Chunk:
    """Représente un morceau de document."""
    
    content: str                    # Contenu textuel
    source_file: str                # Chemin du fichier source
    file_type: str                  # Type de fichier (python, r, markdown...)
    chunk_index: int                # Index du chunk dans le fichier
    total_chunks: int               # Nombre total de chunks du fichier
    start_line: Optional[int] = None   # Ligne de début (si applicable)
    end_line: Optional[int] = None     # Ligne de fin (si applicable)
    section_name: Optional[str] = None # Nom de la section/fonction
    
    @property
    def metadata(self) -> dict:
        """Retourne les métadonnées pour ChromaDB."""
        return {
            "source_file": self.source_file,
            "file_type": self.file_type,
            "chunk_index": self.chunk_index,
            "total_chunks": self.total_chunks,
            "start_line": self.start_line or 0,
            "end_line": self.end_line or 0,
            "section_name": self.section_name or "",
            "char_count": len(self.content),
        }
    
    @property
    def chunk_id(self) -> str:
        """ID unique pour ce chunk."""
        return f"{self.source_file}::{self.chunk_index}"


class BaseChunker(ABC):
    """Classe de base pour les chunkers."""
    
    def __init__(self, config: Optional[ChunkingConfig] = None):
        self.config = config or get_config().chunking
    
    @abstractmethod
    def chunk(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """Découpe le contenu en chunks."""
        pass
    
    def _estimate_tokens(self, text: str) -> int:
        """Estime le nombre de tokens (approximatif : 1 token ≈ 4 caractères)."""
        return len(text) // 4
    
    def _split_with_overlap(self, chunks: List[str], overlap_chars: int) -> List[str]:
        """Ajoute un overlap entre les chunks."""
        if len(chunks) <= 1 or overlap_chars <= 0:
            return chunks
        
        result = []
        for i, chunk in enumerate(chunks):
            if i == 0:
                # Premier chunk : ajouter le début du suivant
                if len(chunks) > 1:
                    next_start = chunks[1][:overlap_chars]
                    chunk = chunk + "\n...\n" + next_start
            elif i == len(chunks) - 1:
                # Dernier chunk : ajouter la fin du précédent
                prev_end = chunks[i-1][-overlap_chars:]
                chunk = prev_end + "\n...\n" + chunk
            else:
                # Chunks du milieu : ajouter des deux côtés
                prev_end = chunks[i-1][-overlap_chars:]
                next_start = chunks[i+1][:overlap_chars] if i+1 < len(chunks) else ""
                chunk = prev_end + "\n...\n" + chunk + "\n...\n" + next_start
            
            result.append(chunk)
        
        return result


class CodeChunker(BaseChunker):
    """Chunker pour le code Python."""
    
    # Patterns pour détecter les structures Python
    FUNCTION_PATTERN = re.compile(
        r'^(async\s+)?def\s+(\w+)\s*\([^)]*\)\s*(?:->.*?)?:',
        re.MULTILINE
    )
    CLASS_PATTERN = re.compile(
        r'^class\s+(\w+)\s*(?:\([^)]*\))?\s*:',
        re.MULTILINE
    )
    SECTION_PATTERN = re.compile(
        r'^# [=]{10,}.*?\n# (.+?)\n# [=]{10,}',
        re.MULTILINE
    )
    
    def chunk(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """Découpe le code Python par fonction/classe/section."""
        chunks = []
        lines = content.split('\n')
        
        # Identifier les positions des fonctions, classes et sections
        boundaries = self._find_boundaries(content)
        
        if not boundaries:
            # Pas de structure détectée : découper par taille
            return self._chunk_by_size(content, source_file, file_type)
        
        # Trier les boundaries par position
        boundaries.sort(key=lambda x: x[0])
        
        # Créer les chunks basés sur les boundaries
        for i, (start_pos, name, block_type) in enumerate(boundaries):
            # Trouver la fin du bloc (début du suivant ou fin du fichier)
            if i + 1 < len(boundaries):
                end_pos = boundaries[i + 1][0]
            else:
                end_pos = len(content)
            
            chunk_content = content[start_pos:end_pos].strip()
            
            if chunk_content and self._estimate_tokens(chunk_content) > 20:
                # Calculer les lignes
                start_line = content[:start_pos].count('\n') + 1
                end_line = start_line + chunk_content.count('\n')
                
                chunks.append(Chunk(
                    content=chunk_content,
                    source_file=source_file,
                    file_type=file_type,
                    chunk_index=len(chunks),
                    total_chunks=0,  # Mis à jour après
                    start_line=start_line,
                    end_line=end_line,
                    section_name=f"{block_type}: {name}"
                ))
        
        # Si les chunks sont trop gros, les subdiviser
        final_chunks = []
        for chunk in chunks:
            if self._estimate_tokens(chunk.content) > self.config.max_chunk_size:
                sub_chunks = self._chunk_by_size(
                    chunk.content, source_file, file_type,
                    base_section=chunk.section_name
                )
                final_chunks.extend(sub_chunks)
            else:
                final_chunks.append(chunk)
        
        # Mettre à jour total_chunks et indices
        for i, chunk in enumerate(final_chunks):
            chunk.chunk_index = i
            chunk.total_chunks = len(final_chunks)
        
        return final_chunks if final_chunks else self._chunk_by_size(content, source_file, file_type)
    
    def _find_boundaries(self, content: str) -> List[Tuple[int, str, str]]:
        """Trouve les positions des fonctions, classes et sections."""
        boundaries = []
        
        # Sections (commentaires ===)
        for match in self.SECTION_PATTERN.finditer(content):
            boundaries.append((match.start(), match.group(1).strip(), "section"))
        
        # Classes
        for match in self.CLASS_PATTERN.finditer(content):
            boundaries.append((match.start(), match.group(1), "class"))
        
        # Fonctions (seulement celles au niveau module, pas dans les classes)
        for match in self.FUNCTION_PATTERN.finditer(content):
            # Vérifier que la fonction n'est pas indentée (niveau module)
            line_start = content.rfind('\n', 0, match.start()) + 1
            indent = match.start() - line_start
            if indent == 0:  # Niveau module
                name = match.group(2)
                boundaries.append((match.start(), name, "function"))
        
        return boundaries
    
    def _chunk_by_size(
        self, 
        content: str, 
        source_file: str, 
        file_type: str,
        base_section: Optional[str] = None
    ) -> List[Chunk]:
        """Découpe par taille avec séparateurs intelligents."""
        max_chars = self.config.max_chunk_size * 4  # Approximation tokens -> chars
        
        # Utiliser les séparateurs de code
        separators = self.config.code_separators
        
        parts = [content]
        for sep in separators:
            new_parts = []
            for part in parts:
                if len(part) > max_chars:
                    split = part.split(sep)
                    new_parts.extend([s + sep for s in split[:-1]] + [split[-1]])
                else:
                    new_parts.append(part)
            parts = new_parts
        
        # Fusionner les parties trop petites
        chunks = []
        current = ""
        for part in parts:
            if len(current) + len(part) < max_chars:
                current += part
            else:
                if current.strip():
                    chunks.append(current.strip())
                current = part
        if current.strip():
            chunks.append(current.strip())
        
        # Créer les objets Chunk
        result = []
        for i, chunk_content in enumerate(chunks):
            section = base_section or f"part_{i+1}"
            result.append(Chunk(
                content=chunk_content,
                source_file=source_file,
                file_type=file_type,
                chunk_index=i,
                total_chunks=len(chunks),
                section_name=section
            ))
        
        return result


class RChunker(BaseChunker):
    """Chunker pour le code R."""
    
    # Patterns pour R
    FUNCTION_PATTERN = re.compile(
        r'^(\w+)\s*<-\s*function\s*\([^)]*\)\s*\{',
        re.MULTILINE
    )
    SECTION_PATTERN = re.compile(
        r'^# [=]{10,}.*?\n# (.+?)\n# [=]{10,}',
        re.MULTILINE
    )
    
    def chunk(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """Découpe le code R par fonction/section."""
        chunks = []
        
        # Identifier les positions des fonctions et sections
        boundaries = []
        
        # Sections
        for match in self.SECTION_PATTERN.finditer(content):
            boundaries.append((match.start(), match.group(1).strip(), "section"))
        
        # Fonctions R
        for match in self.FUNCTION_PATTERN.finditer(content):
            boundaries.append((match.start(), match.group(1), "function"))
        
        if not boundaries:
            return self._chunk_by_size(content, source_file, file_type)
        
        boundaries.sort(key=lambda x: x[0])
        
        for i, (start_pos, name, block_type) in enumerate(boundaries):
            if i + 1 < len(boundaries):
                end_pos = boundaries[i + 1][0]
            else:
                end_pos = len(content)
            
            chunk_content = content[start_pos:end_pos].strip()
            
            if chunk_content and self._estimate_tokens(chunk_content) > 20:
                start_line = content[:start_pos].count('\n') + 1
                end_line = start_line + chunk_content.count('\n')
                
                chunks.append(Chunk(
                    content=chunk_content,
                    source_file=source_file,
                    file_type=file_type,
                    chunk_index=len(chunks),
                    total_chunks=0,
                    start_line=start_line,
                    end_line=end_line,
                    section_name=f"{block_type}: {name}"
                ))
        
        # Subdiviser si nécessaire
        final_chunks = []
        for chunk in chunks:
            if self._estimate_tokens(chunk.content) > self.config.max_chunk_size:
                sub_chunks = self._chunk_by_size(
                    chunk.content, source_file, file_type
                )
                for sc in sub_chunks:
                    sc.section_name = chunk.section_name
                final_chunks.extend(sub_chunks)
            else:
                final_chunks.append(chunk)
        
        for i, chunk in enumerate(final_chunks):
            chunk.chunk_index = i
            chunk.total_chunks = len(final_chunks)
        
        return final_chunks if final_chunks else self._chunk_by_size(content, source_file, file_type)
    
    def _chunk_by_size(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """Découpe par taille pour R."""
        max_chars = self.config.max_chunk_size * 4
        separators = self.config.r_separators
        
        parts = [content]
        for sep in separators:
            new_parts = []
            for part in parts:
                if len(part) > max_chars:
                    split = part.split(sep)
                    new_parts.extend([s + sep for s in split[:-1]] + [split[-1]])
                else:
                    new_parts.append(part)
            parts = new_parts
        
        chunks = []
        current = ""
        for part in parts:
            if len(current) + len(part) < max_chars:
                current += part
            else:
                if current.strip():
                    chunks.append(current.strip())
                current = part
        if current.strip():
            chunks.append(current.strip())
        
        return [
            Chunk(
                content=c,
                source_file=source_file,
                file_type=file_type,
                chunk_index=i,
                total_chunks=len(chunks),
                section_name=f"part_{i+1}"
            )
            for i, c in enumerate(chunks)
        ]


class MarkdownChunker(BaseChunker):
    """Chunker pour les fichiers Markdown."""
    
    HEADER_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
    
    def chunk(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """Découpe le Markdown par sections (headers)."""
        chunks = []
        
        # Trouver tous les headers
        headers = list(self.HEADER_PATTERN.finditer(content))
        
        if not headers:
            return self._chunk_by_paragraphs(content, source_file, file_type)
        
        for i, match in enumerate(headers):
            level = len(match.group(1))
            title = match.group(2).strip()
            start_pos = match.start()
            
            # Fin = prochain header de niveau égal ou supérieur, ou fin
            end_pos = len(content)
            for j in range(i + 1, len(headers)):
                next_level = len(headers[j].group(1))
                if next_level <= level:
                    end_pos = headers[j].start()
                    break
            
            chunk_content = content[start_pos:end_pos].strip()
            
            if chunk_content and self._estimate_tokens(chunk_content) > 20:
                start_line = content[:start_pos].count('\n') + 1
                end_line = start_line + chunk_content.count('\n')
                
                chunks.append(Chunk(
                    content=chunk_content,
                    source_file=source_file,
                    file_type=file_type,
                    chunk_index=len(chunks),
                    total_chunks=0,
                    start_line=start_line,
                    end_line=end_line,
                    section_name=f"h{level}: {title}"
                ))
        
        # Inclure le contenu avant le premier header s'il existe
        if headers and headers[0].start() > 0:
            preamble = content[:headers[0].start()].strip()
            if preamble and self._estimate_tokens(preamble) > 20:
                chunks.insert(0, Chunk(
                    content=preamble,
                    source_file=source_file,
                    file_type=file_type,
                    chunk_index=0,
                    total_chunks=0,
                    section_name="preamble"
                ))
        
        # Subdiviser les gros chunks
        final_chunks = []
        for chunk in chunks:
            if self._estimate_tokens(chunk.content) > self.config.max_chunk_size:
                sub_chunks = self._chunk_by_paragraphs(
                    chunk.content, source_file, file_type
                )
                for sc in sub_chunks:
                    sc.section_name = chunk.section_name
                final_chunks.extend(sub_chunks)
            else:
                final_chunks.append(chunk)
        
        for i, chunk in enumerate(final_chunks):
            chunk.chunk_index = i
            chunk.total_chunks = len(final_chunks)
        
        return final_chunks if final_chunks else self._chunk_by_paragraphs(content, source_file, file_type)
    
    def _chunk_by_paragraphs(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """Découpe par paragraphes."""
        max_chars = self.config.max_chunk_size * 4
        paragraphs = re.split(r'\n\s*\n', content)
        
        chunks = []
        current = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            if len(current) + len(para) + 2 < max_chars:
                current = current + "\n\n" + para if current else para
            else:
                if current:
                    chunks.append(current)
                current = para
        
        if current:
            chunks.append(current)
        
        return [
            Chunk(
                content=c,
                source_file=source_file,
                file_type=file_type,
                chunk_index=i,
                total_chunks=len(chunks),
                section_name=f"paragraph_{i+1}"
            )
            for i, c in enumerate(chunks)
        ]


class TextChunker(BaseChunker):
    """Chunker générique pour le texte."""
    
    def chunk(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """Découpe le texte par paragraphes."""
        max_chars = self.config.max_chunk_size * 4
        
        # Séparer par paragraphes (double saut de ligne)
        paragraphs = re.split(r'\n\s*\n', content)
        
        chunks = []
        current = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            if len(current) + len(para) + 2 < max_chars:
                current = current + "\n\n" + para if current else para
            else:
                if current:
                    chunks.append(current)
                # Si le paragraphe lui-même est trop long, le découper
                if len(para) > max_chars:
                    sentences = re.split(r'(?<=[.!?])\s+', para)
                    current = ""
                    for sent in sentences:
                        if len(current) + len(sent) + 1 < max_chars:
                            current = current + " " + sent if current else sent
                        else:
                            if current:
                                chunks.append(current)
                            current = sent
                else:
                    current = para
        
        if current:
            chunks.append(current)
        
        return [
            Chunk(
                content=c,
                source_file=source_file,
                file_type=file_type,
                chunk_index=i,
                total_chunks=len(chunks),
                section_name=f"segment_{i+1}"
            )
            for i, c in enumerate(chunks)
        ]


class CSVChunker(BaseChunker):
    """Chunker pour les fichiers CSV - garde le header avec chaque chunk."""
    
    def chunk(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """Découpe le CSV en gardant le header."""
        lines = content.strip().split('\n')
        
        if len(lines) < 2:
            return [Chunk(
                content=content,
                source_file=source_file,
                file_type=file_type,
                chunk_index=0,
                total_chunks=1,
                section_name="data"
            )]
        
        header = lines[0]
        data_lines = lines[1:]
        
        # Calculer combien de lignes par chunk
        max_chars = self.config.max_chunk_size * 4
        avg_line_len = sum(len(l) for l in data_lines[:10]) // min(10, len(data_lines))
        lines_per_chunk = max(5, max_chars // (avg_line_len + 1))
        
        chunks = []
        for i in range(0, len(data_lines), lines_per_chunk):
            chunk_lines = data_lines[i:i + lines_per_chunk]
            chunk_content = header + "\n" + "\n".join(chunk_lines)
            
            chunks.append(Chunk(
                content=chunk_content,
                source_file=source_file,
                file_type=file_type,
                chunk_index=len(chunks),
                total_chunks=0,
                start_line=i + 2,  # +2 car header = ligne 1
                end_line=i + len(chunk_lines) + 1,
                section_name=f"rows_{i+1}_to_{i+len(chunk_lines)}"
            ))
        
        for chunk in chunks:
            chunk.total_chunks = len(chunks)
        
        return chunks


class PDFChunker(BaseChunker):
    """
    Chunker pour les fichiers PDF.
    
    Utilise pypdf pour extraire le texte page par page,
    puis découpe intelligemment en respectant les paragraphes.
    """
    
    def chunk(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """
        Découpe le contenu PDF (déjà extrait en texte).
        
        Note: Le contenu est extrait par l'indexer avant d'arriver ici.
        On reçoit donc du texte brut, pas du binaire PDF.
        """
        if not content or len(content.strip()) < 10:
            return []
        
        # Utiliser le TextChunker pour le contenu extrait
        text_chunker = TextChunker(self.config)
        chunks = text_chunker.chunk(content, source_file, file_type)
        
        # Mettre à jour le type pour indiquer l'origine PDF
        for chunk in chunks:
            chunk.file_type = "pdf"
        
        return chunks
    
    @staticmethod
    def extract_text_from_pdf(filepath: str) -> str:
        """
        Extrait le texte d'un fichier PDF.
        
        Args:
            filepath: Chemin vers le fichier PDF
            
        Returns:
            Texte extrait du PDF
        """
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(filepath)
            text_parts = []
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    # Ajouter un marqueur de page pour le contexte
                    text_parts.append(f"[Page {i + 1}]\n{page_text}")
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("pypdf non installé. Installez avec: pip install pypdf")
            return ""
        except Exception as e:
            logger.error(f"Erreur extraction PDF {filepath}: {e}")
            return ""


class ExcelChunker(BaseChunker):
    """
    Chunker pour les fichiers Excel (.xlsx, .xls).
    
    Traite chaque feuille séparément et garde les headers
    avec chaque chunk de données.
    """
    
    def chunk(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """
        Découpe le contenu Excel (déjà converti en texte).
        
        Note: Le contenu est extrait par l'indexer avant d'arriver ici.
        """
        if not content or len(content.strip()) < 10:
            return []
        
        # Si le contenu contient des marqueurs de feuilles, découper par feuille
        if "[Sheet:" in content:
            return self._chunk_by_sheets(content, source_file, file_type)
        
        # Sinon, traiter comme CSV
        csv_chunker = CSVChunker(self.config)
        return csv_chunker.chunk(content, source_file, file_type)
    
    def _chunk_by_sheets(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """Découpe le contenu en séparant les feuilles Excel."""
        import re
        
        # Pattern pour séparer les feuilles : [Sheet: NomFeuille]
        sheet_pattern = re.compile(r'\[Sheet:\s*([^\]]+)\]\n')
        
        # Trouver toutes les feuilles
        sheets = sheet_pattern.split(content)
        
        chunks = []
        i = 1  # Les données commencent à l'index 1 (après le premier match)
        
        while i < len(sheets):
            sheet_name = sheets[i].strip() if i < len(sheets) else f"Sheet_{i//2}"
            sheet_content = sheets[i + 1] if i + 1 < len(sheets) else ""
            
            if sheet_content.strip():
                # Utiliser CSVChunker pour chaque feuille
                csv_chunker = CSVChunker(self.config)
                sheet_chunks = csv_chunker.chunk(sheet_content, source_file, file_type)
                
                # Ajouter le nom de la feuille dans la section
                for chunk in sheet_chunks:
                    chunk.section_name = f"[{sheet_name}] {chunk.section_name}"
                    chunk.file_type = "excel"
                
                chunks.extend(sheet_chunks)
            
            i += 2
        
        # Mettre à jour les indices
        for idx, chunk in enumerate(chunks):
            chunk.chunk_index = idx
            chunk.total_chunks = len(chunks)
        
        return chunks
    
    @staticmethod
    def extract_text_from_excel(filepath: str) -> str:
        """
        Extrait le texte d'un fichier Excel.
        
        Args:
            filepath: Chemin vers le fichier Excel
            
        Returns:
            Texte extrait (format CSV-like avec marqueurs de feuilles)
        """
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(filepath, data_only=True)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                # Extraire les données de la feuille
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    # Convertir les valeurs en strings
                    row_str = [str(cell) if cell is not None else "" for cell in row]
                    # Ignorer les lignes entièrement vides
                    if any(cell.strip() for cell in row_str):
                        rows.append(",".join(row_str))
                
                if rows:
                    text_parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("openpyxl non installé. Installez avec: pip install openpyxl")
            return ""
        except Exception as e:
            logger.error(f"Erreur extraction Excel {filepath}: {e}")
            return ""


class DocxChunker(BaseChunker):
    """
    Chunker pour les fichiers Word (.docx).
    
    Extrait le texte paragraphe par paragraphe et
    préserve la structure du document.
    """
    
    def chunk(self, content: str, source_file: str, file_type: str) -> List[Chunk]:
        """
        Découpe le contenu DOCX (déjà extrait en texte).
        
        Note: Le contenu est extrait par l'indexer avant d'arriver ici.
        """
        if not content or len(content.strip()) < 10:
            return []
        
        # Utiliser MarkdownChunker car la structure est similaire
        # (headers/sections convertis en Markdown-like)
        md_chunker = MarkdownChunker(self.config)
        chunks = md_chunker.chunk(content, source_file, file_type)
        
        # Mettre à jour le type
        for chunk in chunks:
            chunk.file_type = "docx"
        
        return chunks
    
    @staticmethod
    def extract_text_from_docx(filepath: str) -> str:
        """
        Extrait le texte d'un fichier Word (.docx).
        
        Args:
            filepath: Chemin vers le fichier DOCX
            
        Returns:
            Texte extrait avec structure préservée
        """
        try:
            from docx import Document
            
            doc = Document(filepath)
            text_parts = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Détecter les titres par le style
                    style_name = para.style.name.lower() if para.style else ""
                    
                    if "heading 1" in style_name or "titre 1" in style_name:
                        text_parts.append(f"\n# {text}")
                    elif "heading 2" in style_name or "titre 2" in style_name:
                        text_parts.append(f"\n## {text}")
                    elif "heading 3" in style_name or "titre 3" in style_name:
                        text_parts.append(f"\n### {text}")
                    elif "title" in style_name or "titre" in style_name:
                        text_parts.append(f"\n# {text}")
                    else:
                        text_parts.append(text)
            
            # Extraire aussi les tableaux
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("\n[Table]\n" + "\n".join(table_text))
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("python-docx non installé. Installez avec: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"Erreur extraction DOCX {filepath}: {e}")
            return ""


# =============================================================================
# FACTORY
# =============================================================================

def get_chunker(file_type: str) -> BaseChunker:
    """Retourne le chunker approprié pour le type de fichier."""
    chunkers = {
        "python": CodeChunker,
        "r": RChunker,
        "rmarkdown": MarkdownChunker,  # RMarkdown utilise le chunker Markdown
        "markdown": MarkdownChunker,
        "text": TextChunker,
        "csv": CSVChunker,
        "yaml": TextChunker,
        "json": TextChunker,
        "toml": TextChunker,
        "ini": TextChunker,
        "shell": CodeChunker,
        "sql": CodeChunker,
        "javascript": CodeChunker,
        "typescript": CodeChunker,
        "html": TextChunker,
        "css": TextChunker,
        # Nouveaux formats supportés
        "pdf": PDFChunker,
        "excel": ExcelChunker,
        "docx": DocxChunker,
    }
    
    chunker_class = chunkers.get(file_type, TextChunker)
    return chunker_class()


# =============================================================================
# TEST
# =============================================================================

if __name__ == "__main__":
    # Test avec du code Python
    test_code = '''
# ==============================================================================
# MODULE DE TEST
# ==============================================================================

import os
from pathlib import Path

def hello(name: str) -> str:
    """Dit bonjour."""
    return f"Hello, {name}!"

def goodbye(name: str) -> str:
    """Dit au revoir."""
    return f"Goodbye, {name}!"

class Greeter:
    """Classe pour les salutations."""
    
    def __init__(self, name: str):
        self.name = name
    
    def greet(self):
        return hello(self.name)
'''
    
    chunker = CodeChunker()
    chunks = chunker.chunk(test_code, "test.py", "python")
    
    print(f"Nombre de chunks: {len(chunks)}")
    for chunk in chunks:
        print(f"\n--- {chunk.section_name} ---")
        print(chunk.content[:200] + "..." if len(chunk.content) > 200 else chunk.content)
