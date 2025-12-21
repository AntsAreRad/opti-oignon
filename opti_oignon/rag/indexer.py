#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
INDEXER - Indexation des documents dans ChromaDB
=================================================
Parcourt les dossiers, chunke les fichiers et les stocke avec leurs embeddings.

Fonctionnalités :
- Indexation récursive de dossiers
- Mise à jour incrémentale (ne ré-indexe que les modifiés)
- Support de multiples types de fichiers
- Métadonnées riches (source, type, date, etc.)
"""

import hashlib
import json
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple
import fnmatch

import chromadb
from chromadb.config import Settings

from .config import get_config, RAGConfig
from .chunkers import Chunk, get_chunker
from .embeddings import OllamaEmbeddings

logger = logging.getLogger(__name__)


class DocumentIndexer:
    """
    Indexeur de documents pour le système RAG.
    
    Usage:
        indexer = DocumentIndexer()
        stats = indexer.index_directory("/path/to/docs")
        print(f"Indexé: {stats['indexed_files']} fichiers")
    """
    
    def __init__(
        self, 
        config: Optional[RAGConfig] = None,
        collection_name: str = "documents"
    ):
        """
        Initialise l'indexeur.
        
        Args:
            config: Configuration RAG (utilise la config globale si non fourni)
            collection_name: Nom de la collection ChromaDB
        """
        self.config = config or get_config()
        self.collection_name = collection_name
        
        # Initialiser ChromaDB en mode persistant
        self.chroma_client = chromadb.PersistentClient(
            path=str(self.config.chroma_dir),
            settings=Settings(anonymized_telemetry=False)
        )
        
        # Créer ou récupérer la collection
        self.collection = self.chroma_client.get_or_create_collection(
            name=collection_name,
            metadata={"description": "Documents RAG personnels"}
        )
        
        # Embedder
        self.embedder = OllamaEmbeddings(self.config.embedding)
        
        # Cache des fichiers indexés (hash -> metadata)
        self._index_cache_file = self.config.cache_dir / "index_cache.json"
        self._index_cache = self._load_index_cache()
    
    def _load_index_cache(self) -> Dict[str, dict]:
        """Charge le cache d'indexation depuis le disque."""
        if self._index_cache_file.exists():
            try:
                with open(self._index_cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Erreur lecture cache: {e}")
        return {}
    
    def _save_index_cache(self):
        """Sauvegarde le cache d'indexation."""
        try:
            with open(self._index_cache_file, 'w', encoding='utf-8') as f:
                json.dump(self._index_cache, f, indent=2, default=str)
        except Exception as e:
            logger.warning(f"Erreur sauvegarde cache: {e}")
    
    def _compute_file_hash(self, filepath: Path) -> str:
        """Calcule le hash MD5 d'un fichier."""
        hasher = hashlib.md5()
        with open(filepath, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                hasher.update(chunk)
        return hasher.hexdigest()
    
    def _should_skip_file(self, filepath: Path) -> bool:
        """Vérifie si le fichier doit être ignoré."""
        # Vérifier les patterns d'exclusion
        for pattern in self.config.ignore_patterns:
            if fnmatch.fnmatch(filepath.name, pattern):
                return True
            if fnmatch.fnmatch(str(filepath), f"*/{pattern}/*"):
                return True
        
        # Vérifier la taille
        try:
            size_mb = filepath.stat().st_size / (1024 * 1024)
            if size_mb > self.config.max_file_size_mb:
                logger.debug(f"Fichier trop volumineux: {filepath} ({size_mb:.1f} MB)")
                return True
        except Exception:
            return True
        
        return False
    
    def _get_file_type(self, filepath: Path) -> Optional[str]:
        """Retourne le type de fichier basé sur l'extension."""
        ext = filepath.suffix.lower()
        return self.config.file_type_mapping.get(ext)
    
    def _is_file_modified(self, filepath: Path, file_hash: str) -> bool:
        """Vérifie si le fichier a été modifié depuis la dernière indexation."""
        str_path = str(filepath)
        if str_path not in self._index_cache:
            return True
        
        cached = self._index_cache[str_path]
        return cached.get("hash") != file_hash
    
    def _read_file(self, filepath: Path) -> Optional[str]:
        """
        Lit le contenu d'un fichier avec gestion des encodages.
        
        Supporte les formats texte (UTF-8, Latin-1, etc.) et les formats
        binaires (PDF, Excel, DOCX) via extraction spécialisée.
        
        Args:
            filepath: Chemin du fichier à lire
            
        Returns:
            Contenu textuel du fichier, ou None en cas d'erreur
        """
        file_type = self._get_file_type(filepath)
        
        # Traitement spécial pour les formats binaires
        if file_type == "pdf":
            return self._extract_pdf(filepath)
        elif file_type == "excel":
            return self._extract_excel(filepath)
        elif file_type == "docx":
            return self._extract_docx(filepath)
        
        # Formats texte standards
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Erreur lecture {filepath}: {e}")
                return None
        
        logger.warning(f"Impossible de lire {filepath} (encodage inconnu)")
        return None
    
    def _extract_pdf(self, filepath: Path) -> Optional[str]:
        """
        Extrait le texte d'un fichier PDF.
        
        Args:
            filepath: Chemin du fichier PDF
            
        Returns:
            Texte extrait ou None
        """
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(filepath))
            text_parts = []
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {i + 1}]\n{page_text}")
            
            if not text_parts:
                logger.warning(f"PDF sans texte extractible: {filepath}")
                return None
            
            logger.debug(f"PDF extrait: {filepath} ({len(reader.pages)} pages)")
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("pypdf non installé. Installez avec: pip install pypdf")
            return None
        except Exception as e:
            logger.error(f"Erreur extraction PDF {filepath}: {e}")
            return None
    
    def _extract_excel(self, filepath: Path) -> Optional[str]:
        """
        Extrait le texte d'un fichier Excel.
        
        Args:
            filepath: Chemin du fichier Excel
            
        Returns:
            Texte extrait (format CSV-like avec marqueurs de feuilles)
        """
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(str(filepath), data_only=True)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    row_str = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_str):
                        rows.append(",".join(row_str))
                
                if rows:
                    text_parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
            
            if not text_parts:
                logger.warning(f"Excel vide: {filepath}")
                return None
            
            logger.debug(f"Excel extrait: {filepath} ({len(wb.sheetnames)} feuilles)")
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("openpyxl non installé. Installez avec: pip install openpyxl")
            return None
        except Exception as e:
            logger.error(f"Erreur extraction Excel {filepath}: {e}")
            return None
    
    def _extract_docx(self, filepath: Path) -> Optional[str]:
        """
        Extrait le texte d'un fichier Word (.docx).
        
        Args:
            filepath: Chemin du fichier DOCX
            
        Returns:
            Texte extrait avec structure préservée
        """
        try:
            from docx import Document
            
            doc = Document(str(filepath))
            text_parts = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    style_name = para.style.name.lower() if para.style else ""
                    
                    if "heading 1" in style_name or "titre 1" in style_name:
                        text_parts.append(f"\n# {text}")
                    elif "heading 2" in style_name or "titre 2" in style_name:
                        text_parts.append(f"\n## {text}")
                    elif "heading 3" in style_name or "titre 3" in style_name:
                        text_parts.append(f"\n### {text}")
                    elif "title" in style_name or "titre" in style_name:
                        text_parts.append(f"\n# {text}")
                    else:
                        text_parts.append(text)
            
            # Extraire les tableaux
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("\n[Table]\n" + "\n".join(table_text))
            
            if not text_parts:
                logger.warning(f"DOCX vide: {filepath}")
                return None
            
            logger.debug(f"DOCX extrait: {filepath} ({len(doc.paragraphs)} paragraphes)")
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("python-docx non installé. Installez avec: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"Erreur extraction DOCX {filepath}: {e}")
            return None
    
    def index_file(
        self, 
        filepath: Path, 
        force: bool = False
    ) -> Tuple[int, int]:
        """
        Indexe un fichier unique.
        
        Args:
            filepath: Chemin du fichier
            force: Forcer la ré-indexation même si non modifié
            
        Returns:
            Tuple (chunks_added, chunks_skipped)
        """
        filepath = Path(filepath).resolve()
        
        if not filepath.exists():
            logger.warning(f"Fichier non trouvé: {filepath}")
            return (0, 0)
        
        if self._should_skip_file(filepath):
            return (0, 1)
        
        file_type = self._get_file_type(filepath)
        if not file_type:
            logger.debug(f"Type non supporté: {filepath.suffix}")
            return (0, 1)
        
        # Calculer le hash
        file_hash = self._compute_file_hash(filepath)
        
        # Vérifier si modification nécessaire
        if not force and not self._is_file_modified(filepath, file_hash):
            logger.debug(f"Fichier non modifié: {filepath}")
            return (0, 1)
        
        # Supprimer les anciens chunks de ce fichier
        self._remove_file_chunks(filepath)
        
        # Lire le contenu
        content = self._read_file(filepath)
        if not content or len(content.strip()) < 10:
            return (0, 1)
        
        # Chunker le contenu
        chunker = get_chunker(file_type)
        chunks = chunker.chunk(content, str(filepath), file_type)
        
        if not chunks:
            return (0, 1)
        
        # Générer les embeddings
        chunk_texts = [c.content for c in chunks]
        embeddings = self.embedder.embed(chunk_texts, show_progress=False)
        
        # Filtrer les chunks sans embedding valide
        valid_chunks = []
        valid_embeddings = []
        for chunk, emb in zip(chunks, embeddings):
            if emb is not None:
                valid_chunks.append(chunk)
                valid_embeddings.append(emb)
        
        if not valid_chunks:
            logger.warning(f"Aucun embedding valide pour {filepath}")
            return (0, 1)
        
        # Ajouter à ChromaDB
        self.collection.add(
            ids=[c.chunk_id for c in valid_chunks],
            embeddings=valid_embeddings,
            documents=[c.content for c in valid_chunks],
            metadatas=[c.metadata for c in valid_chunks]
        )
        
        # Mettre à jour le cache
        self._index_cache[str(filepath)] = {
            "hash": file_hash,
            "chunks": len(valid_chunks),
            "indexed_at": datetime.now().isoformat(),
            "file_type": file_type
        }
        
        logger.info(f"Indexé: {filepath.name} ({len(valid_chunks)} chunks)")
        return (len(valid_chunks), 0)
    
    def _remove_file_chunks(self, filepath: Path):
        """Supprime tous les chunks d'un fichier."""
        str_path = str(filepath)
        
        # Chercher les chunks existants
        try:
            results = self.collection.get(
                where={"source_file": str_path}
            )
            
            if results and results['ids']:
                self.collection.delete(ids=results['ids'])
                logger.debug(f"Supprimé {len(results['ids'])} chunks de {filepath.name}")
        except Exception as e:
            logger.warning(f"Erreur suppression chunks: {e}")
    
    def index_directory(
        self, 
        directory: Path,
        recursive: bool = True,
        force: bool = False,
        file_types: Optional[List[str]] = None
    ) -> Dict[str, int]:
        """
        Indexe un dossier complet.
        
        Args:
            directory: Chemin du dossier
            recursive: Parcourir les sous-dossiers
            force: Forcer la ré-indexation
            file_types: Liste des types à inclure (None = tous)
            
        Returns:
            Statistiques d'indexation
        """
        directory = Path(directory).resolve()
        
        if not directory.exists():
            raise ValueError(f"Dossier non trouvé: {directory}")
        
        if not directory.is_dir():
            raise ValueError(f"Pas un dossier: {directory}")
        
        stats = {
            "total_files": 0,
            "indexed_files": 0,
            "skipped_files": 0,
            "total_chunks": 0,
            "errors": 0
        }
        
        # Collecter tous les fichiers
        if recursive:
            files = list(directory.rglob("*"))
        else:
            files = list(directory.glob("*"))
        
        # Filtrer les fichiers
        files = [f for f in files if f.is_file()]
        
        # Filtrer par type si spécifié
        if file_types:
            allowed_extensions = set()
            for ft in file_types:
                for ext, type_name in self.config.file_type_mapping.items():
                    if type_name == ft or ext.lstrip('.') == ft:
                        allowed_extensions.add(ext)
            files = [f for f in files if f.suffix.lower() in allowed_extensions]
        
        stats["total_files"] = len(files)
        
        logger.info(f"Indexation de {len(files)} fichiers dans {directory}")
        
        # Indexer chaque fichier
        from tqdm import tqdm
        for filepath in tqdm(files, desc="Indexation", unit="file"):
            try:
                added, skipped = self.index_file(filepath, force=force)
                
                if added > 0:
                    stats["indexed_files"] += 1
                    stats["total_chunks"] += added
                else:
                    stats["skipped_files"] += 1
                    
            except Exception as e:
                logger.error(f"Erreur indexation {filepath}: {e}")
                stats["errors"] += 1
        
        # Sauvegarder le cache
        self._save_index_cache()
        
        logger.info(
            f"Indexation terminée: {stats['indexed_files']} fichiers, "
            f"{stats['total_chunks']} chunks"
        )
        
        return stats
    
    def remove_file(self, filepath: Path) -> bool:
        """
        Supprime un fichier de l'index.
        
        Args:
            filepath: Chemin du fichier
            
        Returns:
            True si supprimé, False sinon
        """
        filepath = Path(filepath).resolve()
        str_path = str(filepath)
        
        try:
            self._remove_file_chunks(filepath)
            
            if str_path in self._index_cache:
                del self._index_cache[str_path]
                self._save_index_cache()
            
            return True
        except Exception as e:
            logger.error(f"Erreur suppression: {e}")
            return False
    
    def clear_index(self) -> bool:
        """
        Vide complètement l'index.
        
        Returns:
            True si succès
        """
        try:
            # Supprimer et recréer la collection
            self.chroma_client.delete_collection(self.collection_name)
            self.collection = self.chroma_client.create_collection(
                name=self.collection_name,
                metadata={"description": "Documents RAG personnels"}
            )
            
            # Vider le cache
            self._index_cache.clear()
            self._save_index_cache()
            
            logger.info("Index vidé avec succès")
            return True
        except Exception as e:
            logger.error(f"Erreur vidage index: {e}")
            return False
    
    def get_stats(self) -> Dict[str, any]:
        """
        Retourne les statistiques de l'index.
        
        Returns:
            Dictionnaire de statistiques
        """
        try:
            count = self.collection.count()
            
            # Compter par type de fichier
            type_counts = {}
            for info in self._index_cache.values():
                ft = info.get("file_type", "unknown")
                type_counts[ft] = type_counts.get(ft, 0) + 1
            
            return {
                "total_chunks": count,
                "total_files": len(self._index_cache),
                "files_by_type": type_counts,
                "collection_name": self.collection_name,
                "storage_path": str(self.config.chroma_dir),
                "embedding_model": self.config.embedding.model
            }
        except Exception as e:
            logger.error(f"Erreur statistiques: {e}")
            return {"error": str(e)}
    
    def list_indexed_files(self) -> List[Dict[str, any]]:
        """
        Liste tous les fichiers indexés.
        
        Returns:
            Liste de métadonnées de fichiers
        """
        files = []
        for filepath, info in self._index_cache.items():
            files.append({
                "path": filepath,
                "type": info.get("file_type"),
                "chunks": info.get("chunks"),
                "indexed_at": info.get("indexed_at")
            })
        
        # Trier par date d'indexation
        files.sort(key=lambda x: x.get("indexed_at", ""), reverse=True)
        return files


# =============================================================================
# FONCTIONS UTILITAIRES
# =============================================================================

def quick_index(directory: str, force: bool = False) -> Dict[str, int]:
    """
    Fonction rapide pour indexer un dossier.
    
    Args:
        directory: Chemin du dossier
        force: Forcer la ré-indexation
        
    Returns:
        Statistiques
    """
    indexer = DocumentIndexer()
    return indexer.index_directory(Path(directory), force=force)


# =============================================================================
# TEST
# =============================================================================

if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO)
    
    indexer = DocumentIndexer()
    
    print("=== Test de l'indexeur ===\n")
    
    # Afficher les stats actuelles
    stats = indexer.get_stats()
    print(f"Fichiers indexés: {stats['total_files']}")
    print(f"Chunks totaux: {stats['total_chunks']}")
    print(f"Modèle: {stats['embedding_model']}")
    
    # Si un argument est passé, indexer ce dossier
    if len(sys.argv) > 1:
        directory = sys.argv[1]
        print(f"\nIndexation de: {directory}")
        result = indexer.index_directory(Path(directory))
        print(f"\nRésultat:")
        for key, value in result.items():
            print(f"  {key}: {value}")
